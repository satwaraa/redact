"""test_audit_redaction — the standalone audit reports leaks and clean runs."""

from __future__ import annotations

import sys
from pathlib import Path

import pytest
from docx import Document

_ROOT = Path(__file__).resolve().parents[1]
if str(_ROOT) not in sys.path:
    sys.path.insert(0, str(_ROOT))

from evaluation.audit_redaction import audit, main, render  # noqa: E402

_SOURCE_LINES = [
    "Contact Person: Priya Deshmukh",
    "Email: priya.deshmukh@northwind-tech.com",
    "Phone: +91 98765 43210",
    "SSN on file 456-78-9012 and card 4111 1111 1111 1111.",
]


def _write(path: Path, lines: list[str]) -> Path:
    doc = Document()
    for line in lines:
        doc.add_paragraph(line)
    # python-docx stamps author="python-docx"; the real pipeline blanks identity
    # fields, and the audit rightly flags them, so clear them in fixtures too.
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.save(str(path))
    return path


def test_audit_flags_an_unredacted_copy(tmp_path: Path) -> None:
    src = _write(tmp_path / "src.docx", _SOURCE_LINES)
    # "Redacted" output that changed nothing at all.
    copy = _write(tmp_path / "copy.docx", _SOURCE_LINES)

    report = audit(src, copy)
    assert report["totals"]["clean"] is False
    assert report["totals"]["applied_leaks"] > 0
    assert report["totals"]["probe_leaks"] > 0
    # The independent probe must see the person the rules also found.
    assert "Priya Deshmukh" in report["probe"]["labelled_person"]["values"]
    assert "LEAKS FOUND" in render(report)


def test_audit_reports_clean_when_values_are_gone(tmp_path: Path) -> None:
    src = _write(tmp_path / "src.docx", _SOURCE_LINES)
    redacted = _write(
        tmp_path / "red.docx",
        [
            "Contact Person: Aryan Maharaj",
            "Email: aryan.maharaj@sample-host.example",
            "Phone: +91 12345 67890",
            "SSN on file 123-45-6780 and card 4111 1111 1111 1111.".replace(
                "4111 1111 1111 1111", "4539 5788 0067 1379"
            ),
        ],
    )
    report = audit(src, redacted)
    assert report["totals"]["applied_leaks"] == 0
    assert report["totals"]["probe_leaks"] == 0
    assert report["totals"]["clean"] is True
    assert "CLEAN" in render(report)


def test_audit_detects_a_package_only_survivor(tmp_path: Path) -> None:
    """A value hidden outside visible text still counts as a leak."""
    src = _write(tmp_path / "src.docx", _SOURCE_LINES)
    redacted = _write(
        tmp_path / "red.docx",
        ["Contact Person: Aryan Maharaj", "Email: someone@sample-host.example"],
    )
    # Smuggle the original address into document metadata.
    doc = Document(str(redacted))
    doc.core_properties.comments = "priya.deshmukh@northwind-tech.com"
    doc.save(str(redacted))

    report = audit(src, redacted)
    email = report["probe"]["email"]
    assert email["survived_package_only"] == 1
    assert email["survived_visible"] == 0
    assert report["totals"]["clean"] is False


def test_main_exit_codes(tmp_path: Path, capsys: pytest.CaptureFixture[str]) -> None:
    src = _write(tmp_path / "src.docx", _SOURCE_LINES)
    copy = _write(tmp_path / "copy.docx", _SOURCE_LINES)
    assert main([str(src), str(copy)]) == 1
    assert main([str(src), str(tmp_path / "missing.docx")]) == 2
    capsys.readouterr()


def test_main_writes_json(tmp_path: Path, capsys: pytest.CaptureFixture[str]) -> None:
    import json

    src = _write(tmp_path / "src.docx", _SOURCE_LINES)
    copy = _write(tmp_path / "copy.docx", _SOURCE_LINES)
    out = tmp_path / "report.json"
    main([str(src), str(copy), "--json", str(out)])
    data = json.loads(out.read_text())
    assert data["totals"]["clean"] is False
    assert "structure" in data
    capsys.readouterr()
