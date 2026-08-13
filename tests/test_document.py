"""test_document — .docx adapter: traversal, offsets, run splicing, save guard."""

from __future__ import annotations

import hashlib
from collections.abc import Callable
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt
from docx.text.paragraph import Paragraph

from pii_redaction.document import DocxDocument
from pii_redaction.models import PRIORITY_REGEX, DocumentError, PIIEntity, PIIType


def _write_docx(path: Path, build: Callable[[Document], None]) -> Path:
    doc = Document()
    # Drop the default empty paragraph but keep sectPr so tables/headers work.
    body = doc.element.body
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)
    build(doc)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path


def _entity(
    text: str,
    start: int,
    *,
    replacement: str | None = "REPLACED",
    pii_type: PIIType = PIIType.EMAIL,
    source: str = "test",
) -> PIIEntity:
    return PIIEntity(
        pii_type=pii_type,
        text=text,
        start=start,
        end=start + len(text),
        source=source,
        priority=PRIORITY_REGEX,
        replacement=replacement,
    )


def _block_texts(doc: DocxDocument) -> list[str]:
    return [b.text for b in doc.blocks]


def _file_digest(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


class TestTraversal:
    def test_body_paragraphs_in_order(self, tmp_path: Path) -> None:
        path = _write_docx(
            tmp_path / "body.docx",
            lambda d: (
                d.add_paragraph("alpha"),
                d.add_paragraph("beta"),
                d.add_paragraph("gamma"),
            ),
        )
        doc = DocxDocument(path)
        assert _block_texts(doc) == ["alpha", "beta", "gamma"]

    def test_table_cells_included(self, tmp_path: Path) -> None:
        def build(d: Document) -> None:
            table = d.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "left"
            table.cell(0, 1).text = "right"

        path = _write_docx(tmp_path / "table.docx", build)
        doc = DocxDocument(path)
        assert "left" in _block_texts(doc)
        assert "right" in _block_texts(doc)

    def test_nested_table_included(self, tmp_path: Path) -> None:
        def build(d: Document) -> None:
            outer = d.add_table(rows=1, cols=1)
            cell = outer.cell(0, 0)
            cell.text = "outer"
            inner = cell.add_table(rows=1, cols=1)
            inner.cell(0, 0).text = "inner"

        path = _write_docx(tmp_path / "nested.docx", build)
        texts = _block_texts(DocxDocument(path))
        assert "outer" in texts
        assert "inner" in texts

    def test_paragraph_table_paragraph_interleaving(self, tmp_path: Path) -> None:
        """Naive paragraphs-then-tables would scramble this order."""

        def build(d: Document) -> None:
            d.add_paragraph("before")
            table = d.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "cell"
            d.add_paragraph("after")

        path = _write_docx(tmp_path / "interleave.docx", build)
        assert _block_texts(DocxDocument(path)) == ["before", "cell", "after"]

    def test_empty_paragraphs_and_zero_run_paragraphs(self, tmp_path: Path) -> None:
        def build(d: Document) -> None:
            d.add_paragraph("")
            d.add_paragraph("kept")
            d.add_paragraph("")

        path = _write_docx(tmp_path / "empty.docx", build)
        assert _block_texts(DocxDocument(path)) == ["", "kept", ""]

    def test_linked_header_appears_once(self, tmp_path: Path) -> None:
        def build(d: Document) -> None:
            d.add_paragraph("body-one")
            d.sections[0].header.paragraphs[0].text = "shared-header"
            d.add_section()
            d.add_paragraph("body-two")
            # section 1 header stays linked to previous → same content once

        path = _write_docx(tmp_path / "header.docx", build)
        texts = _block_texts(DocxDocument(path))
        assert texts.count("shared-header") == 1
        assert "body-one" in texts
        assert "body-two" in texts


class TestOffsetIndex:
    def test_extracted_slices_match_block_text(self, tmp_path: Path) -> None:
        path = _write_docx(
            tmp_path / "offsets.docx",
            lambda d: (d.add_paragraph("one"), d.add_paragraph("two"), d.add_paragraph("")),
        )
        doc = DocxDocument(path)
        extracted = doc.extract_text()
        assert len(doc.blocks) == len(doc._block_spans)  # noqa: SLF001
        for block, (start, end) in zip(doc.blocks, doc._block_spans, strict=True):  # noqa: SLF001
            assert extracted[start:end] == block.text
            assert end - start == len(block.text)

    def test_separator_accounts_for_newline(self, tmp_path: Path) -> None:
        path = _write_docx(
            tmp_path / "sep.docx",
            lambda d: (d.add_paragraph("ab"), d.add_paragraph("cd")),
        )
        doc = DocxDocument(path)
        extracted = doc.extract_text()
        assert extracted == "ab\ncd"
        (_, end0), (start1, _) = doc._block_spans  # noqa: SLF001
        assert start1 == end0 + 1
        assert extracted[end0] == "\n"

    def test_global_offset_resolves_to_block(self, tmp_path: Path) -> None:
        path = _write_docx(
            tmp_path / "resolve.docx",
            lambda d: (d.add_paragraph("aaa"), d.add_paragraph("bbb")),
        )
        doc = DocxDocument(path)
        doc.extract_text()
        idx = doc._block_index_for_offset(4)  # noqa: SLF001  # 'b' of second block
        assert idx == 1
        block_start, _ = doc._block_spans[idx]  # noqa: SLF001
        assert 4 - block_start == 0

    def test_extract_text_stable_across_calls(self, tmp_path: Path) -> None:
        path = _write_docx(
            tmp_path / "stable.docx",
            lambda d: d.add_paragraph("same"),
        )
        doc = DocxDocument(path)
        assert doc.extract_text() == doc.extract_text()


class TestRunSplicing:
    def test_email_split_across_three_runs_round_trip(self, tmp_path: Path) -> None:
        """Acceptance case: email split across three runs; bold survives elsewhere."""

        def build(d: Document) -> None:
            p = d.add_paragraph()
            r0 = p.add_run("rohan.dey")
            r0.bold = True
            r1 = p.add_run("@gmail")
            r1.italic = True
            r2 = p.add_run(".com")
            r2.font.size = Pt(14)
            trail = p.add_run(" OK")
            trail.bold = True

        path = _write_docx(tmp_path / "split.docx", build)
        doc = DocxDocument(path)
        text = doc.extract_text()
        email = "rohan.dey@gmail.com"
        start = text.index(email)
        fake = "alex.kim@example.org"
        doc.apply([_entity(email, start, replacement=fake)])

        paragraph = doc.blocks[0]
        assert paragraph.text == fake + " OK"
        assert email not in paragraph.text
        assert "rohan" not in paragraph.text
        assert "@gmail" not in paragraph.text
        # Interior / covered runs emptied but not deleted; trailing bold run untouched
        assert any(r.text == "" for r in paragraph.runs)
        trailing = [r for r in paragraph.runs if r.text == " OK"]
        assert len(trailing) == 1
        assert trailing[0].bold is True

        out = tmp_path / "split-out.docx"
        doc.save(out)
        reloaded = DocxDocument(out)
        assert fake in reloaded.extract_text()
        assert email not in reloaded.extract_text()
        bold_ok = [
            r
            for b in reloaded.blocks
            for r in b.runs
            if r.text == " OK"
        ]
        assert bold_ok and bold_ok[0].bold is True

    def test_mid_run_to_mid_run_preserves_prefix_and_suffix(self, tmp_path: Path) -> None:
        def build(d: Document) -> None:
            p = d.add_paragraph()
            p.add_run("AA_EMAIL")
            p.add_run("VAL_BB")

        path = _write_docx(tmp_path / "mid.docx", build)
        doc = DocxDocument(path)
        # "EMAILVAL" spans both runs: prefix AA_, suffix _BB
        doc.apply([_entity("EMAILVAL", start=3, replacement="X")])
        assert doc.blocks[0].text == "AA_X_BB"

    def test_span_entirely_within_one_run(self, tmp_path: Path) -> None:
        def build(d: Document) -> None:
            p = d.add_paragraph()
            p.add_run("prefix-secret-suffix")

        path = _write_docx(tmp_path / "within.docx", build)
        doc = DocxDocument(path)
        doc.apply([_entity("secret", start=7, replacement="x")])
        assert doc.blocks[0].text == "prefix-x-suffix"

    def test_span_covers_entire_run(self, tmp_path: Path) -> None:
        def build(d: Document) -> None:
            p = d.add_paragraph()
            p.add_run("keep-")
            p.add_run("GONE")
            p.add_run("-keep")

        path = _write_docx(tmp_path / "whole-run.docx", build)
        doc = DocxDocument(path)
        doc.apply([_entity("GONE", start=5, replacement="NEW")])
        assert doc.blocks[0].text == "keep-NEW-keep"

    def test_replacement_longer_and_shorter(self, tmp_path: Path) -> None:
        def build(d: Document) -> None:
            d.add_paragraph("xxAAyy")
            d.add_paragraph("zzBBww")

        path = _write_docx(tmp_path / "len.docx", build)
        doc = DocxDocument(path)
        text = doc.extract_text()
        e1 = _entity("AA", text.index("AA"), replacement="LONGVAL")
        e2 = _entity("BB", text.index("BB"), replacement="Z")
        doc.apply([e1, e2])
        rendered = doc.rendered_text()
        assert "LONGVAL" in rendered
        assert "Z" in rendered
        assert "AA" not in rendered
        assert "BB" not in rendered

    def test_two_spans_same_paragraph_right_to_left(self, tmp_path: Path) -> None:
        def build(d: Document) -> None:
            d.add_paragraph("one@a.com and two@b.com")

        path = _write_docx(tmp_path / "two.docx", build)
        doc = DocxDocument(path)
        text = doc.extract_text()
        first = "one@a.com"
        second = "two@b.com"
        doc.apply(
            [
                _entity(first, text.index(first), replacement="F1"),
                _entity(second, text.index(second), replacement="F2"),
            ]
        )
        assert doc.blocks[0].text == "F1 and F2"

    def test_emptied_interior_run_still_exists(self, tmp_path: Path) -> None:
        def build(d: Document) -> None:
            p = d.add_paragraph()
            p.add_run("a")
            p.add_run("MID")
            p.add_run("z")

        path = _write_docx(tmp_path / "empty-run.docx", build)
        doc = DocxDocument(path)
        before_count = len(doc.blocks[0].runs)
        # Cover all three runs so the middle is an emptied interior run, not deleted.
        doc.apply([_entity("aMIDz", start=0, replacement="X")])
        runs = list(doc.blocks[0].runs)
        assert len(runs) == before_count
        assert runs[1].text == ""
        assert doc.blocks[0].text == "X"


class TestPreconditionsAndFailure:
    def test_overlapping_entities_raise(self, tmp_path: Path) -> None:
        path = _write_docx(
            tmp_path / "overlap.docx",
            lambda d: d.add_paragraph("abcdefgh"),
        )
        doc = DocxDocument(path)
        with pytest.raises(DocumentError, match="overlapping"):
            doc.apply(
                [
                    _entity("abcd", 0, replacement="1"),
                    _entity("cdef", 2, replacement="2"),
                ]
            )

    def test_non_docx_raises_document_error(self, tmp_path: Path) -> None:
        bad = tmp_path / "not.docx"
        bad.write_text("not a docx", encoding="utf-8")
        with pytest.raises(DocumentError, match="cannot open"):
            DocxDocument(bad)

    def test_missing_file_raises_document_error(self, tmp_path: Path) -> None:
        missing = tmp_path / "missing.docx"
        with pytest.raises(DocumentError, match=str(missing.resolve())):
            DocxDocument(missing)

    def test_save_refuses_resolved_input_path(self, tmp_path: Path) -> None:
        path = _write_docx(
            tmp_path / "data" / "x.docx",
            lambda d: d.add_paragraph("hi"),
        )
        doc = DocxDocument(path)
        with pytest.raises(DocumentError, match="refusing to overwrite"):
            doc.save(path)
        # Textually different, same resolved path
        with pytest.raises(DocumentError, match="refusing to overwrite"):
            doc.save(tmp_path / "data" / ".." / "data" / "x.docx")

    def test_multi_block_entity_rejected(self, tmp_path: Path) -> None:
        path = _write_docx(
            tmp_path / "multi.docx",
            lambda d: (d.add_paragraph("aaa"), d.add_paragraph("bbb")),
        )
        doc = DocxDocument(path)
        # Fabricate an entity that claims to span the newline + both sides
        bogus = PIIEntity(
            pii_type=PIIType.EMAIL,
            text="aa\nbb",
            start=1,
            end=6,
            source="test",
            priority=PRIORITY_REGEX,
            replacement="X",
        )
        with pytest.raises(DocumentError, match="multiple blocks"):
            doc.apply([bogus])


class TestRoundTrip:
    def test_load_apply_save_reload(self, tmp_path: Path) -> None:
        def build(d: Document) -> None:
            p = d.add_paragraph()
            r = p.add_run("visible ")
            r.italic = True
            p.add_run("secret@mail.com")

        src = _write_docx(tmp_path / "in.docx", build)
        before = _file_digest(src)
        doc = DocxDocument(src)
        original = doc.extract_text()
        email = "secret@mail.com"
        start = original.index(email)
        doc.apply([_entity(email, start, replacement="fake@ex.com")])
        assert doc.rendered_text() != original
        assert "fake@ex.com" in doc.rendered_text()

        out = tmp_path / "out.docx"
        doc.save(out)
        assert _file_digest(src) == before

        reloaded = DocxDocument(out)
        assert "fake@ex.com" in reloaded.extract_text()
        assert email not in reloaded.extract_text()
        italics = [r for b in reloaded.blocks for r in b.runs if r.text == "visible "]
        assert italics and italics[0].italic is True

    def test_rendered_text_differs_from_pre_apply_extract(self, tmp_path: Path) -> None:
        path = _write_docx(
            tmp_path / "cache.docx",
            lambda d: d.add_paragraph("keep leak@x.com end"),
        )
        doc = DocxDocument(path)
        before = doc.extract_text()
        email = "leak@x.com"
        doc.apply([_entity(email, before.index(email), replacement="ok@y.com")])
        assert doc.rendered_text() != before
        assert "ok@y.com" in doc.rendered_text()
        # Cache invalidated: extract_text reflects new content
        assert doc.extract_text() == doc.rendered_text()


class TestDocumentedLimitations:
    def test_value_spanning_two_paragraphs_not_one_entity(self, tmp_path: Path) -> None:
        path = _write_docx(
            tmp_path / "span.docx",
            lambda d: (d.add_paragraph("user@"), d.add_paragraph("mail.com")),
        )
        doc = DocxDocument(path)
        extracted = doc.extract_text()
        assert "user@\nmail.com" in extracted
        # Each block is separate; no single block contains the full address
        assert not any("user@mail.com" in b.text for b in doc.blocks)

    def test_blocks_are_paragraphs(self, tmp_path: Path) -> None:
        path = _write_docx(
            tmp_path / "types.docx",
            lambda d: d.add_paragraph("only"),
        )
        doc = DocxDocument(path)
        assert all(isinstance(b, Paragraph) for b in doc.blocks)
