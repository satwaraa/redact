"""test_redactor — full pipeline: resolve, assign, apply, verify, logging."""

from __future__ import annotations

import logging
import zipfile
from pathlib import Path

import pytest
from docx import Document

from pii_redaction.document import DocxDocument, package_corpus
from pii_redaction.models import (
    PRIORITY_REGEX,
    LeakDetectedError,
    PIIEntity,
    PIIType,
    RedactorConfig,
)
from pii_redaction.redactor import (
    Redactor,
    apply_text,
    expand_occurrences,
    verify_no_leaks,
    verify_package_no_leaks,
    verify_rule_recall,
)


def _write_simple_docx(path: Path, text: str) -> Path:
    doc = Document()
    body = doc.element.body
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)
    doc.add_paragraph(text)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path


def _write_docx(path: Path, paragraphs: list[str]) -> Path:
    doc = Document()
    body = doc.element.body
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path


def _cfg(**kwargs: object) -> RedactorConfig:
    base = {"use_ner": False, "seed": 0, "verify_output": True}
    base.update(kwargs)
    return RedactorConfig(**base)  # type: ignore[arg-type]


def test_redact_text_replaces_pii_and_keeps_non_pii() -> None:
    text = "Contact rashhi.patil@gmail.com or +91 9876543210 today"
    result = Redactor(_cfg()).redact_text(text)
    assert result.text is not None
    assert "rashhi.patil@gmail.com" not in result.text
    assert "+91 9876543210" not in result.text
    assert "Contact " in result.text
    assert " today" in result.text
    assert result.entities
    assert all(e.replacement for e in result.entities)
    assert dict(result.mapping)
    assert PIIType.EMAIL in result.counts_by_type
    assert PIIType.PHONE in result.counts_by_type


def test_assignment_example_regression() -> None:
    text = "Rashi Patil / rashhi.patil@gmail.com / +91 9876543210"
    result = Redactor(_cfg()).redact_text(text)
    assert result.text is not None
    assert "rashhi.patil@gmail.com" not in result.text
    assert "+91 9876543210" not in result.text


def test_no_pii_unchanged() -> None:
    for sample in ("", "   ", "hello world", "no pii here"):
        result = Redactor(_cfg()).redact_text(sample)
        assert result.text == sample
        assert result.entities == ()
        assert dict(result.mapping) == {}


def test_text_and_docx_paths_agree(tmp_path: Path) -> None:
    text = "mail a@b.co and call +91 9876543210 please"
    redactor = Redactor(_cfg(seed=42))
    text_result = redactor.redact_text(text)

    src = _write_simple_docx(tmp_path / "in.docx", text)
    dst = tmp_path / "out.docx"
    doc_result = Redactor(_cfg(seed=42)).redact_document(src, dst)

    assert text_result.text == DocxDocument(dst).extract_text()
    assert {e.pii_type for e in text_result.entities} == {
        e.pii_type for e in doc_result.entities
    }


def test_document_dry_run_writes_nothing(tmp_path: Path) -> None:
    src = _write_simple_docx(tmp_path / "in.docx", "a@b.co")
    dst = tmp_path / "out.docx"
    result = Redactor(_cfg()).redact_document(src, dst, dry_run=True)
    assert not dst.exists()
    assert any(e.pii_type is PIIType.EMAIL for e in result.entities)
    assert all(e.replacement for e in result.entities)


def test_leak_when_apply_is_noop(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    src = _write_simple_docx(tmp_path / "in.docx", "secret@mail.com lives here")
    dst = tmp_path / "out.docx"

    def _noop_apply(self: DocxDocument, entities: object) -> None:  # noqa: ARG001
        return None

    monkeypatch.setattr(DocxDocument, "apply", _noop_apply)
    with pytest.raises(LeakDetectedError) as exc_info:
        Redactor(_cfg()).redact_document(src, dst)
    assert "EMAIL" in str(exc_info.value)
    assert "secret@mail.com" not in str(exc_info.value)
    assert not dst.exists()


def test_verify_output_false_skips_leak_check(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    src = _write_simple_docx(tmp_path / "in.docx", "secret@mail.com")
    dst = tmp_path / "out.docx"
    monkeypatch.setattr(DocxDocument, "apply", lambda self, entities: None)
    result = Redactor(_cfg(verify_output=False)).redact_document(src, dst)
    assert dst.exists()
    assert result.entities


def test_short_surname_false_positive_guard() -> None:
    # "Ann" is below the leak length threshold and may survive inside "Anniversary"
    entities = [
        PIIEntity(
            pii_type=PIIType.FULL_NAME,
            text="Ann",
            start=0,
            end=3,
            source="test",
            priority=PRIORITY_REGEX,
            replacement="Zoe",
        )
    ]
    rendered = "Anniversary party"
    verify_no_leaks(entities, rendered, original_text=None)


def test_expand_occurrences_covers_missed_repeats() -> None:
    text = "Acme Technologies Pvt Ltd won; later Acme Technologies Pvt Ltd lost"
    first = "Acme Technologies Pvt Ltd"
    entities = [
        PIIEntity(
            pii_type=PIIType.COMPANY,
            text=first,
            start=0,
            end=len(first),
            source="ner:spacy",
            priority=PRIORITY_REGEX,
        )
    ]
    expanded = expand_occurrences(text, entities)
    assert len(expanded) == 2
    assert expanded[0].start == 0
    assert expanded[1].start == text.rfind(first)
    assert expanded[1].text == first


def test_verify_ignores_original_substring_inside_surrogate() -> None:
    # Faker may embed a short real token inside an unrelated fake; offset-level
    # checks still require the source span itself to change.
    original = "Alpha"
    entities = [
        PIIEntity(
            pii_type=PIIType.COMPANY,
            text=original,
            start=0,
            end=5,
            source="test",
            priority=PRIORITY_REGEX,
            replacement="AlphaBeta Corp",
        )
    ]
    rendered = "AlphaBeta Corp nearby"
    verify_no_leaks(entities, rendered, original_text=original)


def test_cross_block_longer_span_does_not_leave_in_block_repeat(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """Reproduce the prospectus COMPANY pattern without the full doc.

    A longer detection spans two paragraphs (dropped before apply). A shorter
    company string appears once under that span and once alone. After the
    cross-block drop, the lone repeat must still be replaced or D7 fails.
    """
    company = "Acme Technologies Pvt Ltd"
    src = _write_docx(
        tmp_path / "in.docx",
        [company + " is", " Headquarters listed", company],
    )

    flat = DocxDocument(src).extract_text()
    # Longer span starts at the company prefix and crosses into block 1.
    long_start = flat.find(company)
    long_end = flat.find("Headquarters") + len("Headquarters")
    assert "\n" in flat[long_start:long_end]
    longer_text = flat[long_start:long_end]
    alone_start = flat.rfind(company)
    assert alone_start > long_start

    injected = [
        PIIEntity(
            pii_type=PIIType.COMPANY,
            text=longer_text,
            start=long_start,
            end=long_end,
            source="ner:spacy",
            priority=PRIORITY_REGEX,
        ),
        PIIEntity(
            pii_type=PIIType.COMPANY,
            text=company,
            start=alone_start,
            end=alone_start + len(company),
            source="ner:spacy",
            priority=PRIORITY_REGEX,
        ),
    ]

    redactor = Redactor(_cfg(use_ner=False, seed=0))
    monkeypatch.setattr(redactor, "detect", lambda _text: injected)

    dst = tmp_path / "out.docx"
    result = redactor.redact_document(src, dst)
    assert dst.exists()
    out_text = DocxDocument(dst).extract_text()
    assert company not in out_text
    assert any(e.pii_type is PIIType.COMPANY for e in result.entities)


def test_determinism_two_runs(tmp_path: Path) -> None:
    text = "email a@b.co phone +91 9876543210"
    a = Redactor(_cfg(seed=7)).redact_text(text)
    b = Redactor(_cfg(seed=7)).redact_text(text)
    assert a.text == b.text
    assert [e.replacement for e in a.entities] == [e.replacement for e in b.entities]


def test_logging_never_contains_pii_values(caplog: pytest.LogCaptureFixture) -> None:
    text = "reach rashhi.patil@gmail.com or +91 9876543210"
    with caplog.at_level(logging.DEBUG, logger="pii_redaction"):
        Redactor(_cfg()).redact_text(text)
    blob = " ".join(rec.getMessage() for rec in caplog.records)
    assert "rashhi.patil@gmail.com" not in blob
    assert "+91 9876543210" not in blob
    assert "9876543210" not in blob


def test_apply_text_right_to_left() -> None:
    text = "aa@b.co and cc@d.co"
    e1 = PIIEntity(
        pii_type=PIIType.EMAIL,
        text="aa@b.co",
        start=0,
        end=7,
        source="t",
        priority=PRIORITY_REGEX,
        replacement="ONE",
    )
    e2 = PIIEntity(
        pii_type=PIIType.EMAIL,
        text="cc@d.co",
        start=12,
        end=19,
        source="t",
        priority=PRIORITY_REGEX,
        replacement="TWO",
    )
    assert apply_text(text, [e1, e2]) == "ONE and TWO"


def test_verify_rule_recall_flags_unreplaced_value() -> None:
    with pytest.raises(LeakDetectedError) as exc_info:
        verify_rule_recall(
            "contact manisha.shukla@hdfcbank.com please",
            "still manisha.shukla@hdfcbank.com here",
            _cfg(),
        )
    assert "EMAIL" in str(exc_info.value)


def test_verify_rule_recall_passes_when_replaced() -> None:
    verify_rule_recall(
        "contact manisha.shukla@hdfcbank.com please",
        "contact replaced@example.com please",
        _cfg(),
    )


def test_verify_rule_recall_flags_unreplaced_phone() -> None:
    with pytest.raises(LeakDetectedError) as exc_info:
        verify_rule_recall(
            "call +91 9876543210 now",
            "still +91 9876543210 here",
            _cfg(),
        )
    assert "PHONE" in str(exc_info.value)


def test_c2_does_not_flag_phone_shaped_xml_attribute_ids(tmp_path: Path) -> None:
    """C2 runs on story text; OOXML attribute digit runs must not abort redaction."""
    src = _write_simple_docx(tmp_path / "in.docx", "plain prospectus prose")
    with zipfile.ZipFile(src, "r") as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
        other = {
            name: zf.read(name) for name in zf.namelist() if name != "word/document.xml"
        }
    marker = "</w:body>"
    assert marker in xml
    # 10-digit w:id values match the bare phone regex on raw package_corpus.
    snippet = (
        '<w:bookmarkStart w:id="9876543210" w:name="noise"/>'
        '<w:bookmarkEnd w:id="9876543210"/>'
    )
    patched = xml.replace(marker, snippet + marker, 1)
    tmp = src.with_suffix(".tmp.docx")
    with zipfile.ZipFile(tmp, "w") as zf:
        zf.writestr("word/document.xml", patched.encode("utf-8"))
        for name, data in other.items():
            zf.writestr(name, data)
    tmp.replace(src)

    assert "9876543210" in package_corpus(src)
    dst = tmp_path / "out.docx"
    Redactor(_cfg()).redact_document(src, dst)
    assert dst.exists()
    assert "9876543210" in package_corpus(dst)


def test_cross_line_digit_runs_are_not_phones(tmp_path: Path) -> None:
    """Table/year fragments across paragraphs must not become C2 phone leaks."""
    src = _write_docx(
        tmp_path / "in.docx",
        ["December 2024", "1", "234", "56", "Employee count", "100", "200", "300", "400"],
    )
    dst = tmp_path / "out.docx"
    Redactor(_cfg()).redact_document(src, dst)
    assert dst.exists()


def test_verify_package_no_leaks_scans_raw_parts(tmp_path: Path) -> None:
    email = "hidden@mail.com"
    src = _write_simple_docx(tmp_path / "pkg.docx", "no visible pii")
    _inject_instr_mailto(src, email)
    entities = [
        PIIEntity(
            pii_type=PIIType.EMAIL,
            text=email,
            start=0,
            end=len(email),
            source="test",
            priority=PRIORITY_REGEX,
            replacement="x@y.co",
        )
    ]
    with pytest.raises(LeakDetectedError) as exc_info:
        verify_package_no_leaks(entities, src)
    assert "EMAIL" in str(exc_info.value)


def test_field_code_email_is_redacted_from_package(tmp_path: Path) -> None:
    """A1: mailto inside instrText is detectable and removed from the package."""
    email = "manisha.shukla@hdfcbank.com"
    src = _write_simple_docx(tmp_path / "in.docx", "visible a@b.co only")
    _inject_instr_mailto(src, email)
    dst = tmp_path / "out.docx"
    result = Redactor(_cfg()).redact_document(src, dst)
    assert dst.exists()
    assert email not in package_corpus(dst)
    assert any(e.pii_type is PIIType.EMAIL for e in result.entities)


def _inject_instr_mailto(path: Path, email: str) -> None:
    """Append a HYPERLINK mailto field into word/document.xml (extractor-blind)."""
    with zipfile.ZipFile(path, "r") as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
        other = {
            name: zf.read(name) for name in zf.namelist() if name != "word/document.xml"
        }

    marker = "</w:body>"
    assert marker in xml
    snippet = (
        "<w:p><w:r>"
        f'<w:instrText xml:space="preserve"> HYPERLINK "mailto:{email}" </w:instrText>'
        "</w:r></w:p>"
    )
    patched = xml.replace(marker, snippet + marker, 1)
    tmp = path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(tmp, "w") as zf:
        zf.writestr("word/document.xml", patched.encode("utf-8"))
        for name, data in other.items():
            zf.writestr(name, data)
    tmp.replace(path)
